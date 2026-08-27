"""Comprehensive tests for exports.py, export_engine.py, and listing_template.py."""

import hashlib
import os
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


def _make_idea(**overrides):
    idea = {
        "name": "Premium Blender",
        "asin": "B0TEST001",
        "category": "Kitchen",
        "amazon_price": 29.99,
        "price": 29.99,
        "rating": 4.5,
        "review_count": 1200,
        "estimated_margin_pct": 38.5,
        "score": 0.72,
        "gated": False,
        "gated_category": "",
        "url": "https://amazon.com/dp/B0TEST001",
        "image": "https://img.example.com/1.jpg",
        "images": ["https://img.example.com/1.jpg", "https://img.example.com/2.jpg"],
        "priority": {"rank": 1, "tier": "CRITICAL", "action": "BUY NOW"},
        "marketing": {
            "marketing_score": 0.81,
            "summary": "Strong demand in kitchen.",
            "problems": [
                {"severity": "HIGH", "problem": "Competition", "description": "Many sellers."},
                {"severity": "LOW", "problem": "Reviews", "description": "Few negative."},
            ],
            "solutions": [
                {
                    "solution": "Bundle Offer",
                    "priority": "High",
                    "estimated_cost": "£150",
                    "timeline": "2 weeks",
                    "actions": ["Create listing", "Set price"],
                },
            ],
            "recommended_strategies": [
                {"name": "PPC Campaign", "priority": "High", "reason": "Visibility", "cost": "£200", "time_to_results": "1 week"},
            ],
        },
        "listing_template": {
            "product_identity": {
                "item_name": "Premium Blender Pro",
                "brand_name": "BlendCo",
                "product_type": "Blender",
            },
            "description": {
                "bullet_points": ["Powerful motor", "Easy clean"],
                "product_description": "A high-quality blender.",
                "images": ["https://img.example.com/1.jpg"],
            },
            "product_details": {
                "material": "Plastic",
                "colour": "Red",
                "size": "Large",
            },
            "offer": {
                "sku": "PRE-1234ABCD",
                "your_price": "29.99",
                "item_condition": "New",
            },
            "_metadata": {
                "estimated_margin_pct": 38.5,
            },
        },
    }
    idea.update(overrides)
    return idea


def _mock_docx_objects():
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    mock_doc = MagicMock()
    mock_heading = MagicMock()
    mock_heading.alignment = None
    mock_doc.add_heading.return_value = mock_heading
    mock_para = MagicMock()
    mock_para.alignment = None
    mock_run = MagicMock()
    mock_run.font = MagicMock(size=None, color=MagicMock(rgb=None))
    mock_para.add_run.return_value = mock_run
    mock_doc.add_paragraph.return_value = mock_para
    mock_table = MagicMock()
    mock_table.alignment = None
    mock_table.rows = [
        MagicMock(cells=[
            MagicMock(text="", paragraphs=[MagicMock(runs=[], style=MagicMock(font=MagicMock(size=None)))])
            for _ in range(4)
        ])
        for _ in range(20)
    ]
    mock_doc.add_table.return_value = mock_table
    return mock_doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT


def _patch_word_module(mod, mock_doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT):
    return [
        patch.object(mod, "HAS_DOCX", True),
        patch.object(mod, "Document", return_value=mock_doc),
        patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH),
        patch.object(mod, "Pt", Pt),
        patch.object(mod, "RGBColor", RGBColor),
        patch.object(mod, "Inches", MagicMock()),
        patch.object(mod, "WD_TABLE_ALIGNMENT", WD_TABLE_ALIGNMENT),
    ]


# ===================================================================
# PART 1 - tests for utils/exports.py
# ===================================================================


class TestExportToExcelImportError:
    def test_raises_when_openpyxl_missing(self):
        import utils.exports as mod
        with patch.object(mod, "HAS_OPENPYXL", False):
            with pytest.raises(ImportError, match="openpyxl not installed"):
                mod.export_to_excel([], "/tmp")


class TestExportToExcel:
    def test_creates_three_sheets(self, tmp_path):
        import utils.exports as mod
        mock_wb = MagicMock()
        mock_wb.active = MagicMock()
        mock_wb.create_sheet.return_value = MagicMock()
        with patch.object(mod, "Workbook", return_value=mock_wb):
            ideas = [_make_idea(), _make_idea(name="Another", asin="B002")]
            result = mod.export_to_excel(ideas, str(tmp_path), "test.xlsx")
            assert result.endswith("test.xlsx")
            mock_wb.create_sheet.assert_any_call("Marketing Analysis")
            mock_wb.create_sheet.assert_any_call("Listing Templates")
            mock_wb.save.assert_called_once()

    def test_empty_ideas(self, tmp_path):
        import utils.exports as mod
        mock_wb = MagicMock()
        mock_wb.active = MagicMock()
        mock_wb.create_sheet.return_value = MagicMock()
        with patch.object(mod, "Workbook", return_value=mock_wb):
            result = mod.export_to_excel([], str(tmp_path))
            mock_wb.save.assert_called_once()

    def test_ideas_missing_optional_fields(self, tmp_path):
        import utils.exports as mod
        mock_wb = MagicMock()
        mock_wb.active = MagicMock()
        mock_wb.create_sheet.return_value = MagicMock()
        with patch.object(mod, "Workbook", return_value=mock_wb):
            mod.export_to_excel([{"name": "Minimal"}], str(tmp_path), "min.xlsx")
            mock_wb.save.assert_called_once()

    def test_priority_tiers_all_branches(self, tmp_path):
        import utils.exports as mod
        for tier in ["CRITICAL", "HIGH", "MEDIUM", "LOW", "MINIMAL"]:
            mock_wb = MagicMock()
            mock_wb.active = MagicMock()
            mock_wb.create_sheet.return_value = MagicMock()
            with patch.object(mod, "Workbook", return_value=mock_wb):
                idea = _make_idea(priority={"rank": 1, "tier": tier, "action": "Test"})
                mod.export_to_excel([idea], str(tmp_path), f"tier_{tier}.xlsx")
                mock_wb.save.assert_called_once()


class TestWriteIdeasSheet:
    def test_writes_header_and_data(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_ideas_sheet(mock_ws, [_make_idea()])
        assert mock_ws.cell.call_count > 0

    def test_hyperlink_on_image_column(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_ideas_sheet(mock_ws, [_make_idea(url="https://example.com")])
        assert mock_ws.cell.call_count > 0

    def test_gated_product(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_ideas_sheet(mock_ws, [_make_idea(gated=True, gated_category="Kitchen")])

    def test_ungated_product(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_ideas_sheet(mock_ws, [_make_idea(gated=False)])

    def test_price_fallback(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        idea = _make_idea()
        idea["amazon_price"] = None
        mod._write_ideas_sheet(mock_ws, [idea])

    def test_tier_unknown(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_ideas_sheet(mock_ws, [_make_idea(priority={"rank": 1, "tier": "UNKNOWN", "action": ""})])

    def test_missing_priority(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        idea = _make_idea()
        del idea["priority"]
        mod._write_ideas_sheet(mock_ws, [idea])


class TestWriteMarketingSheet:
    def test_full_marketing(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_marketing_sheet(mock_ws, [_make_idea()])

    def test_empty_marketing(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_marketing_sheet(mock_ws, [_make_idea(marketing={})])

    def test_empty_problems_solutions_strategies(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        marketing = {"marketing_score": 0.5, "summary": "T", "problems": [], "solutions": [], "recommended_strategies": []}
        mod._write_marketing_sheet(mock_ws, [_make_idea(marketing=marketing)])

    def test_multiple_problems(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        marketing = {
            "marketing_score": 0.6, "summary": "ok",
            "problems": [
                {"severity": "HIGH", "problem": "A", "description": "Desc A"},
                {"severity": "MEDIUM", "problem": "B", "description": "Desc B"},
            ],
            "solutions": [], "recommended_strategies": [],
        }
        mod._write_marketing_sheet(mock_ws, [_make_idea(marketing=marketing)])

    def test_solutions_with_many_actions(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        marketing = {
            "marketing_score": 0.7, "summary": "ok", "problems": [],
            "solutions": [{"solution": "Fix A", "priority": "High", "actions": ["a1", "a2", "a3", "a4"]}],
            "recommended_strategies": [],
        }
        mod._write_marketing_sheet(mock_ws, [_make_idea(marketing=marketing)])

    def test_multiple_strategies(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        marketing = {
            "marketing_score": 0.8, "summary": "ok", "problems": [], "solutions": [],
            "recommended_strategies": [
                {"name": "PPC", "priority": "High", "reason": "Vis"},
                {"name": "Social", "priority": "Low", "reason": "Brand"},
            ],
        }
        mod._write_marketing_sheet(mock_ws, [_make_idea(marketing=marketing)])


class TestWriteListingsSheet:
    def test_full_listing(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_listings_sheet(mock_ws, [_make_idea()])

    def test_empty_listing_template(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        mod._write_listings_sheet(mock_ws, [_make_idea(listing_template={})])

    def test_empty_bullet_points(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        idea = _make_idea()
        idea["listing_template"]["description"]["bullet_points"] = ["", "valid"]
        mod._write_listings_sheet(mock_ws, [idea])

    def test_no_images(self):
        import utils.exports as mod
        mock_ws = MagicMock()
        idea = _make_idea()
        idea["listing_template"]["description"]["images"] = []
        mod._write_listings_sheet(mock_ws, [idea])


class TestExportToWordImportError:
    def test_raises_when_docx_missing(self):
        import utils.exports as mod
        with patch.object(mod, "HAS_DOCX", False):
            with pytest.raises(ImportError, match="python-docx not installed"):
                mod.export_to_word({}, "/tmp")


class TestExportToWord:
    def test_creates_doc_with_filename(self, tmp_path):
        import utils.exports as mod
        mock_doc, Pt, RGBColor, WDA, WDT = _mock_docx_objects()
        patches = _patch_word_module(mod, mock_doc, Pt, RGBColor, WDA, WDT)
        with patches[0], patches[1], patches[2], patches[3], patches[4], patches[5], patches[6]:
            result = mod.export_to_word(_make_idea(), str(tmp_path), "test.docx")
            assert result.endswith("test.docx")
            mock_doc.save.assert_called_once()

    def test_auto_generates_filename(self, tmp_path):
        import utils.exports as mod
        mock_doc, Pt, RGBColor, WDA, WDT = _mock_docx_objects()
        patches = _patch_word_module(mod, mock_doc, Pt, RGBColor, WDA, WDT)
        with patches[0], patches[1], patches[2], patches[3], patches[4], patches[5], patches[6]:
            result = mod.export_to_word(_make_idea(), str(tmp_path))
            assert result.endswith(".docx")
            safe = "".join(c if c.isalnum() else "_" for c in "Premium Blender")
            assert safe[:50] in os.path.basename(result)

    def test_special_chars_in_name(self, tmp_path):
        import utils.exports as mod
        mock_doc, Pt, RGBColor, WDA, WDT = _mock_docx_objects()
        patches = _patch_word_module(mod, mock_doc, Pt, RGBColor, WDA, WDT)
        with patches[0], patches[1], patches[2], patches[3], patches[4], patches[5], patches[6]:
            result = mod.export_to_word(_make_idea(name="My Product!!! @#$%"), str(tmp_path))
            stem = os.path.basename(result).replace(".docx", "")
            assert all(c.isalnum() or c == "_" for c in stem)


class TestAddTitle:
    def test_priority_critical(self):
        import utils.exports as mod
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_doc, _, _, _, _ = _mock_docx_objects()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "RGBColor", RGBColor), patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH):
            mod._add_title(mock_doc, _make_idea(priority={"rank": 1, "tier": "CRITICAL", "action": "Go"}))

    def test_priority_high(self):
        import utils.exports as mod
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_doc, _, _, _, _ = _mock_docx_objects()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "RGBColor", RGBColor), patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH):
            mod._add_title(mock_doc, _make_idea(priority={"rank": 1, "tier": "HIGH", "action": "Go"}))

    def test_priority_medium(self):
        import utils.exports as mod
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_doc, _, _, _, _ = _mock_docx_objects()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "RGBColor", RGBColor), patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH):
            mod._add_title(mock_doc, _make_idea(priority={"rank": 1, "tier": "MEDIUM", "action": "Go"}))

    def test_priority_low(self):
        import utils.exports as mod
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_doc, _, _, _, _ = _mock_docx_objects()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "RGBColor", RGBColor), patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH):
            mod._add_title(mock_doc, _make_idea(priority={"rank": 1, "tier": "LOW", "action": "Go"}))

    def test_no_priority(self):
        import utils.exports as mod
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_doc, _, _, _, _ = _mock_docx_objects()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "RGBColor", RGBColor), patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH):
            mod._add_title(mock_doc, _make_idea(priority={}))

    def test_no_action(self):
        import utils.exports as mod
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_doc, _, _, _, _ = _mock_docx_objects()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "RGBColor", RGBColor), patch.object(mod, "WD_ALIGN_PARAGRAPH", WD_ALIGN_PARAGRAPH):
            mod._add_title(mock_doc, _make_idea(priority={"rank": 1, "tier": "MEDIUM", "action": ""}))


class TestAddProductIdentity:
    def _make_mock_doc_table(self):
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_table.rows = [
            MagicMock(cells=[
                MagicMock(text="", paragraphs=[MagicMock(runs=[], style=MagicMock(font=MagicMock(size=None)))])
                for _ in range(2)
            ])
            for _ in range(6)
        ]
        mock_doc.add_table.return_value = mock_table
        return mock_doc

    def test_with_listing_template(self):
        import utils.exports as mod
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        mock_doc = self._make_mock_doc_table()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "WD_TABLE_ALIGNMENT", WD_TABLE_ALIGNMENT):
            mod._add_product_identity(mock_doc, _make_idea())
            mock_doc.add_heading.assert_called_once()

    def test_without_listing_template(self):
        import utils.exports as mod
        from docx.shared import Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        mock_doc = self._make_mock_doc_table()
        with patch.object(mod, "Pt", Pt), patch.object(mod, "WD_TABLE_ALIGNMENT", WD_TABLE_ALIGNMENT):
            mod._add_product_identity(mock_doc, _make_idea(listing_template={}))


class TestAddListingContent:
    def test_with_description(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        mod._add_listing_content(mock_doc, _make_idea())
        assert mock_doc.add_heading.call_count >= 2

    def test_empty_description(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        idea = _make_idea()
        idea["listing_template"]["description"]["product_description"] = ""
        mod._add_listing_content(mock_doc, idea)

    def test_empty_images(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        idea = _make_idea()
        idea["listing_template"]["description"]["images"] = []
        mod._add_listing_content(mock_doc, idea)

    def test_empty_bullet_points(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        idea = _make_idea()
        idea["listing_template"]["description"]["bullet_points"] = []
        mod._add_listing_content(mock_doc, idea)

    def test_no_listing_template(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        mod._add_listing_content(mock_doc, _make_idea(listing_template={}))

    def test_many_images_truncated(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        idea = _make_idea()
        idea["listing_template"]["description"]["images"] = [f"https://img.example.com/{i}.jpg" for i in range(10)]
        mod._add_listing_content(mock_doc, idea)


class TestAddMarketingAnalysis:
    def test_full_marketing(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        with patch.object(mod, "RGBColor", RGBColor):
            mod._add_marketing_analysis(mock_doc, _make_idea())
            assert mock_doc.add_heading.call_count >= 3

    def test_no_marketing(self):
        import utils.exports as mod
        mock_doc = MagicMock()
        mod._add_marketing_analysis(mock_doc, _make_idea(marketing={}))
        mock_doc.add_paragraph.assert_any_call("No marketing analysis available.")

    def test_no_problems(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        with patch.object(mod, "RGBColor", RGBColor):
            m = {"marketing_score": 0.5, "summary": "ok", "problems": [], "solutions": [], "recommended_strategies": []}
            mod._add_marketing_analysis(mock_doc, _make_idea(marketing=m))

    def test_no_solutions(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        with patch.object(mod, "RGBColor", RGBColor):
            m = {"marketing_score": 0.5, "summary": "ok",
                 "problems": [{"severity": "HIGH", "problem": "X", "description": "Y"}],
                 "solutions": [], "recommended_strategies": []}
            mod._add_marketing_analysis(mock_doc, _make_idea(marketing=m))

    def test_no_strategies(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        with patch.object(mod, "RGBColor", RGBColor):
            m = {"marketing_score": 0.5, "summary": "ok", "problems": [],
                 "solutions": [{"solution": "Fix", "priority": "High", "estimated_cost": "£10", "timeline": "1d", "actions": ["Do"]}],
                 "recommended_strategies": []}
            mod._add_marketing_analysis(mock_doc, _make_idea(marketing=m))

    def test_empty_summary(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        with patch.object(mod, "RGBColor", RGBColor):
            mod._add_marketing_analysis(mock_doc, _make_idea(marketing={"marketing_score": 0.5, "summary": ""}))

    def test_problem_severity_not_high(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        with patch.object(mod, "RGBColor", RGBColor):
            m = {"marketing_score": 0.5, "summary": "",
                 "problems": [{"severity": "LOW", "problem": "Minor", "description": "D"}],
                 "solutions": [], "recommended_strategies": []}
            mod._add_marketing_analysis(mock_doc, _make_idea(marketing=m))

    def test_strategies_table(self):
        import utils.exports as mod
        from docx.shared import RGBColor
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_table.rows = [
            MagicMock(cells=[MagicMock(text="", paragraphs=[MagicMock(runs=[MagicMock(bold=False)])]) for _ in range(4)])
            for _ in range(20)
        ]
        mock_doc.add_table.return_value = mock_table
        with patch.object(mod, "RGBColor", RGBColor):
            m = {"marketing_score": 0.5, "summary": "ok", "problems": [], "solutions": [],
                 "recommended_strategies": [
                     {"name": "PPC", "priority": "High", "cost": "£200", "time_to_results": "1 week"},
                     {"name": "Social", "priority": "Low", "cost": "£50", "time_to_results": "1 month"},
                 ]}
            mod._add_marketing_analysis(mock_doc, _make_idea(marketing=m))


class TestAddMetadata:
    def _make_mock_doc_table(self):
        mock_doc = MagicMock()
        mock_table = MagicMock()
        mock_table.rows = [
            MagicMock(cells=[MagicMock(text="", paragraphs=[MagicMock(runs=[MagicMock(bold=False)])]) for _ in range(2)])
            for _ in range(10)
        ]
        mock_doc.add_table.return_value = mock_table
        return mock_doc

    def test_with_full_listing(self):
        import utils.exports as mod
        from docx.enum.table import WD_TABLE_ALIGNMENT
        mock_doc = self._make_mock_doc_table()
        with patch.object(mod, "WD_TABLE_ALIGNMENT", WD_TABLE_ALIGNMENT):
            mod._add_metadata(mock_doc, _make_idea())

    def test_gated_category(self):
        import utils.exports as mod
        from docx.enum.table import WD_TABLE_ALIGNMENT
        mock_doc = self._make_mock_doc_table()
        with patch.object(mod, "WD_TABLE_ALIGNMENT", WD_TABLE_ALIGNMENT):
            mod._add_metadata(mock_doc, _make_idea(gated=True, gated_category="Kitchen"))

    def test_not_gated(self):
        import utils.exports as mod
        from docx.enum.table import WD_TABLE_ALIGNMENT
        mock_doc = self._make_mock_doc_table()
        with patch.object(mod, "WD_TABLE_ALIGNMENT", WD_TABLE_ALIGNMENT):
            mod._add_metadata(mock_doc, _make_idea(gated=False))


class TestExportAllToWord:
    def _patch(self, mod):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        mock_doc, _, _, _, _ = _mock_docx_objects()
        return _patch_word_module(mod, mock_doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT)

    def test_multiple_ideas(self, tmp_path):
        import utils.exports as mod
        ps = self._patch(mod)
        with ps[0], ps[1], ps[2], ps[3], ps[4], ps[5], ps[6]:
            result = mod.export_all_to_word([_make_idea(name="A"), _make_idea(name="B")], str(tmp_path))
            assert len(result) == 2
            assert all(p.endswith(".docx") for p in result)

    def test_empty_ideas(self, tmp_path):
        import utils.exports as mod
        with patch.object(mod, "HAS_DOCX", True):
            result = mod.export_all_to_word([], str(tmp_path))
            assert result == []

    def test_creates_directory(self, tmp_path):
        import utils.exports as mod
        nested = tmp_path / "sub" / "dir"
        ps = self._patch(mod)
        with ps[0], ps[1], ps[2], ps[3], ps[4], ps[5], ps[6]:
            result = mod.export_all_to_word([_make_idea()], str(nested))
            assert len(result) == 1
            assert nested.exists()

    def test_filename_indexing(self, tmp_path):
        import utils.exports as mod
        ps = self._patch(mod)
        with ps[0], ps[1], ps[2], ps[3], ps[4], ps[5], ps[6]:
            ideas = [_make_idea(name=f"Prod{i}") for i in range(3)]
            result = mod.export_all_to_word(ideas, str(tmp_path))
            for i, path in enumerate(result, 1):
                assert os.path.basename(path).startswith(f"{i:03d}_")


# ===================================================================
# PART 2 - tests for utils/export_engine.py
# ===================================================================


class TestExcelExporter:
    def test_export_products_creates_file(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "out.xlsx")
        exporter = ExcelExporter()
        products = [
            {"asin": "B01", "name": "A", "ai_score": 0.9, "estimated_margin_pct": 45, "traffic_light": "GREEN"},
            {"asin": "B02", "name": "B", "ai_score": 0.6, "estimated_margin_pct": 20, "traffic_light": "RED"},
        ]
        result = exporter.export_products(products, filepath)
        assert result == filepath
        assert os.path.isfile(filepath)

    def test_creates_core_sheets(self, tmp_path):
        from openpyxl import load_workbook
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "sheets.xlsx")
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5}], filepath)
        wb = load_workbook(filepath)
        names = wb.sheetnames
        assert "Products" in names
        assert "Profitability" in names
        assert "AI Analysis" in names
        assert "Consistency" in names
        assert "Forecasting" in names
        assert "Marketing" in names
        assert "Seller Info" in names
        assert "Supplier Details" in names
        assert "Summary" in names
        assert len(names) >= 9

    def test_empty_products(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "empty.xlsx")
        exporter = ExcelExporter()
        exporter.export_products([], filepath)
        assert os.path.isfile(filepath)

    def test_margin_coloring_green(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "green.xlsx")
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "estimated_margin_pct": 50, "traffic_light": "GREEN"}], filepath)
        assert os.path.isfile(filepath)

    def test_margin_coloring_yellow(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "yellow.xlsx")
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "estimated_margin_pct": 30, "traffic_light": "YELLOW"}], filepath)
        assert os.path.isfile(filepath)

    def test_margin_coloring_red(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "red.xlsx")
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "estimated_margin_pct": 10, "traffic_light": "RED"}], filepath)
        assert os.path.isfile(filepath)

    def test_traffic_light_colors(self, tmp_path):
        from utils.export_engine import ExcelExporter
        for color in ["GREEN", "YELLOW", "RED"]:
            filepath = str(tmp_path / f"tl_{color}.xlsx")
            exporter = ExcelExporter()
            exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "estimated_margin_pct": 30, "traffic_light": color}], filepath)
            assert os.path.isfile(filepath)

    def test_priority_as_dict(self, tmp_path):
        from utils.export_engine import ExcelExporter
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "priority": {"tier": "HIGH"}}], str(tmp_path / "d.xlsx"))
        assert os.path.isfile(str(tmp_path / "d.xlsx"))

    def test_priority_not_dict(self, tmp_path):
        from utils.export_engine import ExcelExporter
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "priority": "HIGH"}], str(tmp_path / "s.xlsx"))
        assert os.path.isfile(str(tmp_path / "s.xlsx"))

    def test_name_fallback_to_title(self, tmp_path):
        from utils.export_engine import ExcelExporter
        exporter = ExcelExporter()
        exporter.export_products([{"asin": "B01", "title": "Title Only", "ai_score": 0.5}], str(tmp_path / "t.xlsx"))
        assert os.path.isfile(str(tmp_path / "t.xlsx"))

    def test_seller_info_sheet(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "si.xlsx")
        exporter = ExcelExporter()
        products = [{"asin": "B01", "name": "X", "ai_score": 0.5,
                      "seller_info": {"seller_name": "Bob", "is_fba": True, "is_prime": True,
                                      "seller_rating": 4.5, "seller_reviews": 100, "brand": "BrandX",
                                      "bsr": 1000, "monthly_sales_est": 500, "competition_level": "Medium",
                                      "is_amazon_retail": False, "seller_location": "US"}}]
        exporter.export_products(products, filepath)
        assert os.path.isfile(filepath)

    def test_seller_info_fbm(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "fbm.xlsx")
        exporter = ExcelExporter()
        products = [{"asin": "B01", "name": "X", "ai_score": 0.5,
                      "seller_info": {"seller_name": "Bob", "is_fba": False, "is_prime": False}}]
        exporter.export_products(products, filepath)
        assert os.path.isfile(filepath)

    def test_summary_sheet_counts(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "sum.xlsx")
        exporter = ExcelExporter()
        products = [
            {"asin": "B01", "name": "X", "ai_score": 0.8, "estimated_margin_pct": 45, "traffic_light": "GREEN"},
            {"asin": "B02", "name": "Y", "ai_score": 0.3, "estimated_margin_pct": 10, "traffic_light": "RED"},
        ]
        exporter.export_products(products, filepath)
        assert os.path.isfile(filepath)

    def test_supplier_pricing_with_suppliers(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "sp.xlsx")
        mock_supplier_db = MagicMock()
        mock_supplier_db.get_suppliers_for_category.return_value = [{"name": "SupplierA", "contact_email": "a@b.com"}]
        mock_pricing = MagicMock()
        mock_pricing.generate_pricing.return_value = {
            "unit_cost": 5.0, "moq": 100, "lead_time_days": 14,
            "shipping_cost_per_unit": 1.5, "customs_duty": 0.5,
            "total_landed_cost": 7.0, "fba_fee": 4.0, "profit_per_unit": 18.99,
            "margin_percent": 40.0,
        }
        mock_module = MagicMock(SupplierDatabase=mock_supplier_db, SupplierPricing=mock_pricing)
        with patch.dict("sys.modules", {"data_collectors.supplier_intel": mock_module}):
            exporter = ExcelExporter()
            products = [{"asin": "B01", "name": "X", "ai_score": 0.5, "category": "Kitchen", "amazon_price": 29.99}]
            exporter.export_products(products, filepath)
            assert os.path.isfile(filepath)

    def test_supplier_pricing_no_suppliers(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "ns.xlsx")
        mock_supplier_db = MagicMock()
        mock_supplier_db.get_suppliers_for_category.return_value = []
        mock_pricing = MagicMock()
        mock_module = MagicMock(SupplierDatabase=mock_supplier_db, SupplierPricing=mock_pricing)
        with patch.dict("sys.modules", {"data_collectors.supplier_intel": mock_module}):
            exporter = ExcelExporter()
            exporter.export_products([{"asin": "B01", "name": "X", "ai_score": 0.5, "category": "Kitchen", "amazon_price": 29.99}], filepath)
            assert os.path.isfile(filepath)

    def test_many_products_sorted_by_score(self, tmp_path):
        from utils.export_engine import ExcelExporter
        filepath = str(tmp_path / "sort.xlsx")
        exporter = ExcelExporter()
        products = [
            {"asin": f"B{i:03d}", "name": f"P{i}", "ai_score": 0.1 * i}
            for i in range(5, 0, -1)
        ]
        exporter.export_products(products, filepath)
        from openpyxl import load_workbook
        wb = load_workbook(filepath)
        ws = wb["Products"]
        first_name = ws.cell(row=2, column=2).value
        assert "P5" in str(first_name)


class TestPDFExporter:
    def test_export_report(self, tmp_path):
        from utils.export_engine import PDFExporter
        filepath = str(tmp_path / "report.pdf")
        exporter = PDFExporter()
        products = [
            {"asin": "B01", "name": "A", "ai_score": 0.9, "estimated_margin_pct": 45,
             "traffic_light": "GREEN", "category": "Kitchen", "amazon_price": 29.99},
        ]
        result = exporter.export_report(products, filepath)
        assert result == filepath
        assert os.path.isfile(filepath)

    def test_export_report_empty(self, tmp_path):
        from utils.export_engine import PDFExporter
        filepath = str(tmp_path / "empty.pdf")
        exporter = PDFExporter()
        exporter.export_report([], filepath)
        assert os.path.isfile(filepath)

    def test_create_summary_section(self):
        from utils.export_engine import PDFExporter
        exporter = PDFExporter()
        products = [
            {"asin": "B01", "name": "A", "ai_score": 0.8, "estimated_margin_pct": 50, "traffic_light": "GREEN"},
            {"asin": "B02", "name": "B", "ai_score": 0.3, "estimated_margin_pct": 10, "traffic_light": "RED"},
        ]
        elements = exporter._create_executive_summary(products)
        assert len(elements) > 0

    def test_create_top_products_section(self):
        from utils.export_engine import PDFExporter
        exporter = PDFExporter()
        products = [
            {"asin": "B01", "name": "A", "ai_score": 0.9, "estimated_margin_pct": 45,
             "traffic_light": "GREEN", "category": "Kitchen", "amazon_price": 29.99},
            {"asin": "B02", "name": "B", "ai_score": 0.6, "estimated_margin_pct": 20,
             "traffic_light": "RED", "category": "Beauty", "amazon_price": 14.99},
        ]
        elements = exporter._create_top_products_section(products)
        assert len(elements) > 0

    def test_create_category_section(self):
        from utils.export_engine import PDFExporter
        exporter = PDFExporter()
        products = [
            {"asin": "B01", "name": "A", "ai_score": 0.8, "estimated_margin_pct": 40,
             "category": "Kitchen", "amazon_price": 29.99},
            {"asin": "B02", "name": "B", "ai_score": 0.6, "estimated_margin_pct": 20,
             "category": "Kitchen", "amazon_price": 19.99},
            {"asin": "B03", "name": "C", "ai_score": 0.7, "estimated_margin_pct": 30,
             "category": "Beauty", "amazon_price": 24.99},
        ]
        elements = exporter._create_category_section(products)
        assert len(elements) > 0

    def test_top_products_sorted_by_ai(self):
        from utils.export_engine import PDFExporter
        exporter = PDFExporter()
        products = [
            {"asin": "B01", "name": "Low", "ai_score": 0.3, "estimated_margin_pct": 20,
             "traffic_light": "RED", "category": "X", "amazon_price": 10.00},
            {"asin": "B02", "name": "High", "ai_score": 0.9, "estimated_margin_pct": 50,
             "traffic_light": "GREEN", "category": "X", "amazon_price": 50.00},
        ]
        elements = exporter._create_top_products_section(products)
        assert len(elements) > 0

    def test_title_override(self, tmp_path):
        from utils.export_engine import PDFExporter
        filepath = str(tmp_path / "custom.pdf")
        exporter = PDFExporter()
        exporter.export_report([], filepath, title="Custom Title")
        assert os.path.isfile(filepath)

    def test_category_section_empty(self):
        from utils.export_engine import PDFExporter
        exporter = PDFExporter()
        elements = exporter._create_category_section([])
        assert len(elements) > 0


# ===================================================================
# PART 3 - tests for utils/listing_template.py
# ===================================================================


class TestGenerateSku:
    def test_normal_name(self):
        from utils.listing_template import _generate_sku
        sku = _generate_sku("Premium Blender Pro")
        assert sku.startswith("PBP-")
        assert len(sku) == 12

    def test_empty_name(self):
        from utils.listing_template import _generate_sku
        assert _generate_sku("") == ""

    def test_single_word(self):
        from utils.listing_template import _generate_sku
        sku = _generate_sku("Blender")
        assert sku.startswith("B-")
        assert len(sku) == 10

    def test_deterministic(self):
        from utils.listing_template import _generate_sku
        assert _generate_sku("Test Product") == _generate_sku("Test Product")

    def test_different_names_different_skus(self):
        from utils.listing_template import _generate_sku
        assert _generate_sku("Product A") != _generate_sku("Product B")

    def test_special_characters(self):
        from utils.listing_template import _generate_sku
        sku = _generate_sku("My Product!!! @#$")
        assert "-" in sku

    def test_two_words(self):
        from utils.listing_template import _generate_sku
        sku = _generate_sku("Test Product")
        assert sku.startswith("TP-")


class TestGenerateListingTemplate:
    def test_with_seo(self):
        from utils.listing_template import generate_listing_template
        mock_seo_class = MagicMock()
        mock_seo_class.return_value.optimize_listing.return_value = {
            "optimized_title": "Optimized Title",
            "optimized_bullets": ["Bullet 1", "Bullet 2"],
            "search_terms": "kw1 kw2",
            "seo_analysis": {
                "primary_keywords": ["kw1"],
                "long_tail_keywords": ["long kw1"],
                "backend_keywords": ["backend1"],
                "seo_score": {"overall": 85},
                "optimization_tips": ["Tip 1"],
            },
        }
        with patch("utils.listing_template.SEOAnalyzer", mock_seo_class):
            template = generate_listing_template(_make_idea(), include_seo=True)
            assert template["product_identity"]["external_product_id"] == "B0TEST001"
            assert template["product_identity"]["item_name"] == "Optimized Title"
            assert template["description"]["bullet_points"] == ["Bullet 1", "Bullet 2"]
            assert template["seo"]["primary_keywords"] == ["kw1"]
            assert template["seo"]["search_terms"] == "kw1 kw2"

    def test_without_seo(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(), include_seo=False)
        assert template["product_identity"]["item_name"] == "Premium Blender"
        assert template["description"]["bullet_points"] == ["", "", "", "", ""]

    def test_seo_exception_fallback(self):
        from utils.listing_template import generate_listing_template
        with patch("utils.listing_template.SEOAnalyzer", side_effect=Exception("fail")):
            template = generate_listing_template(_make_idea(), include_seo=True)
            assert template["product_identity"]["item_name"] == "Premium Blender"

    def test_images_from_image_field(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(images=[], image="https://img.example.com/single.jpg"), include_seo=False)
        assert template["description"]["images"] == ["https://img.example.com/single.jpg"]

    def test_no_images_at_all(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(images=[], image=""), include_seo=False)
        assert template["description"]["images"] == []

    def test_price_zero(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(amazon_price=0), include_seo=False)
        assert template["offer"]["your_price"] == ""

    def test_price_with_value(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(amazon_price=29.99), include_seo=False)
        assert template["offer"]["your_price"] == "29.99"

    def test_metadata_fields(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(score=0.72, tier="CRITICAL"), include_seo=False)
        assert template["_metadata"]["score"] == 0.72
        assert template["_metadata"]["tier"] == "CRITICAL"
        assert template["_metadata"]["generated_from"] == "amazon-product-ai"

    def test_gated_product(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(gated=True, gated_category="Kitchen"), include_seo=False)
        assert template["_metadata"]["is_gated"] is True
        assert template["_metadata"]["gated_category"] == "Kitchen"

    def test_minimal_idea(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template({}, include_seo=False)
        assert template["product_identity"]["item_name"] == ""
        assert template["offer"]["sku"] == ""
        assert template["_metadata"]["asin_reference"] == ""

    def test_reference_image_from_images(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(images=["https://img.example.com/first.jpg"]), include_seo=False)
        assert template["_metadata"]["reference_image"] == "https://img.example.com/first.jpg"

    def test_empty_images_reference_empty(self):
        from utils.listing_template import generate_listing_template
        template = generate_listing_template(_make_idea(images=[], image=""), include_seo=False)
        assert template["_metadata"]["reference_image"] == ""

    def test_seo_analysis_passthrough(self):
        from utils.listing_template import generate_listing_template
        mock_seo_class = MagicMock()
        mock_seo_class.return_value.optimize_listing.return_value = {
            "optimized_title": "Title",
            "optimized_bullets": ["B1"],
            "search_terms": "terms",
            "seo_analysis": {
                "primary_keywords": [],
                "long_tail_keywords": [],
                "backend_keywords": [],
                "seo_score": {},
                "optimization_tips": [],
            },
        }
        with patch("utils.listing_template.SEOAnalyzer", mock_seo_class):
            template = generate_listing_template(_make_idea(), include_seo=True)
            assert template["seo"]["primary_keywords"] == []
            assert template["seo"]["seo_score"] == {}


class TestListingTemplateToText:
    def test_full_template(self):
        from utils.listing_template import listing_template_to_text, generate_listing_template
        template = generate_listing_template(_make_idea(), include_seo=False)
        text = listing_template_to_text(template)
        assert "AMAZON LISTING TEMPLATE" in text
        assert "PRODUCT IDENTITY" in text
        assert "DESCRIPTION" in text
        assert "PRODUCT DETAILS" in text
        assert "OFFER" in text
        assert "PACKAGE DIMENSIONS" in text
        assert "SAFETY AND COMPLIANCE" in text

    def test_gated_metadata(self):
        from utils.listing_template import listing_template_to_text
        template = {"_metadata": {"asin_reference": "B01", "source_url": "", "market_rating": 4.5,
                                   "market_reviews": 100, "score": 0.7, "tier": "HIGH",
                                   "estimated_margin_pct": 35, "is_gated": True, "gated_category": "Kitchen"}}
        text = listing_template_to_text(template)
        assert "GATED" in text
        assert "Kitchen" in text

    def test_not_gated(self):
        from utils.listing_template import listing_template_to_text
        template = {"_metadata": {"asin_reference": "B01", "source_url": "", "market_rating": 4.5,
                                   "market_reviews": 100, "score": 0.7, "tier": "HIGH",
                                   "estimated_margin_pct": 35, "is_gated": False}}
        text = listing_template_to_text(template)
        assert "GATED" not in text

    def test_empty_template(self):
        from utils.listing_template import listing_template_to_text
        text = listing_template_to_text({})
        assert "AMAZON LISTING TEMPLATE" in text

    def test_list_fields(self):
        from utils.listing_template import listing_template_to_text
        template = {"description": {"bullet_points": ["Powerful", "Easy"], "images": []}, "_metadata": {}}
        text = listing_template_to_text(template)
        assert "Powerful" in text
        assert "Easy" in text

    def test_empty_list_item(self):
        from utils.listing_template import listing_template_to_text
        template = {"description": {"bullet_points": ["", "Valid"]}, "_metadata": {}}
        text = listing_template_to_text(template)
        assert "(empty)" in text

    def test_filled_vs_empty_values(self):
        from utils.listing_template import listing_template_to_text
        template = {"product_identity": {"item_name": "Test", "brand_name": ""}, "_metadata": {}}
        text = listing_template_to_text(template)
        assert "Test" in text
        assert "(to be filled)" in text

    def test_empty_section_skipped(self):
        from utils.listing_template import listing_template_to_text
        text = listing_template_to_text({"product_identity": {}, "_metadata": {}})
        assert "PRODUCT IDENTITY" not in text
