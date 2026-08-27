"""Export utilities for Excel and Word formats."""

import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def export_to_excel(
    ideas: List[Dict[str, Any]],
    output_path: str,
    filename: str = "product_ideas.xlsx",
) -> str:
    """Export product ideas to Excel spreadsheet."""
    if not HAS_OPENPYXL:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")

    wb = Workbook()

    ws_ideas = wb.active
    ws_ideas.title = "Product Ideas"
    _write_ideas_sheet(ws_ideas, ideas)

    ws_marketing = wb.create_sheet("Marketing Analysis")
    _write_marketing_sheet(ws_marketing, ideas)

    ws_listings = wb.create_sheet("Listing Templates")
    _write_listings_sheet(ws_listings, ideas)

    filepath = os.path.join(output_path, filename)
    wb.save(filepath)
    return filepath


def _write_ideas_sheet(ws, ideas: List[Dict[str, Any]]):
    """Write product ideas to worksheet."""
    headers = [
        "Priority Rank", "Priority Tier", "Product Name", "ASIN", "Category", "Price (£)",
        "Rating", "Reviews", "Margin (%)", "Access",
        "Score", "Marketing Score", "Action", "URL", "Image URL"
    ]

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    tier_fills = {
        "CRITICAL": PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid"),
        "HIGH": PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid"),
        "MEDIUM": PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid"),
        "LOW": PatternFill(start_color="0000FF", end_color="0000FF", fill_type="solid"),
        "MINIMAL": PatternFill(start_color="808080", end_color="808080", fill_type="solid"),
    }

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = thin_border

    for row_idx, idea in enumerate(ideas, 2):
        priority = idea.get("priority", {})
        priority_tier = priority.get("tier", "unknown")
        marketing = idea.get("marketing", {})

        data = [
            priority.get("rank", row_idx - 1),
            priority_tier,
            idea.get("name", ""),
            idea.get("asin", ""),
            idea.get("category", ""),
            idea.get("amazon_price", 0) or idea.get("price", 0),
            idea.get("rating", 0),
            idea.get("review_count", 0),
            idea.get("estimated_margin_pct", 0),
            "Gated" if idea.get("gated") else "Ungated",
            idea.get("score", 0),
            marketing.get("marketing_score", 0),
            priority.get("action", ""),
            idea.get("url", ""),
            idea.get("image", ""),
        ]

        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.border = thin_border
            if col in [6, 7, 8, 9, 11, 12]:
                cell.number_format = "0.00"
            if col == 15 and value:
                cell.hyperlink = value
                cell.font = Font(color="0563C1", underline="single")

        tier_cell = ws.cell(row=row_idx, column=2)
        if priority_tier in tier_fills:
            tier_cell.fill = tier_fills[priority_tier]
            if priority_tier in ["CRITICAL", "HIGH"]:
                tier_cell.font = Font(bold=True, color="FFFFFF")
            else:
                tier_cell.font = Font(bold=True)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18

    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(ideas) + 1}"


def _write_marketing_sheet(ws, ideas: List[Dict[str, Any]]):
    """Write marketing analysis to worksheet."""
    headers = [
        "Product Name", "Marketing Score", "Summary",
        "Problems", "Solutions", "Recommended Strategies"
    ]

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="548235", end_color="548235", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = thin_border

    for row_idx, idea in enumerate(ideas, 2):
        marketing = idea.get("marketing", {})
        problems = marketing.get("problems", [])
        solutions = marketing.get("solutions", [])
        strategies = marketing.get("recommended_strategies", [])

        problems_text = "\n".join([
            f"- [{p.get('severity', '').upper()}] {p.get('problem', '')}: {p.get('description', '')}"
            for p in problems
        ])

        solutions_text = "\n".join([
            f"- {s.get('solution', '')} (Priority: {s.get('priority', '')})\n  Actions: {'; '.join(s.get('actions', [])[:3])}"
            for s in solutions
        ])

        strategies_text = "\n".join([
            f"- {st.get('name', '')} ({st.get('priority', '')}) - {st.get('reason', '')}"
            for st in strategies
        ])

        data = [
            idea.get("name", ""),
            marketing.get("marketing_score", 0),
            marketing.get("summary", ""),
            problems_text,
            solutions_text,
            strategies_text,
        ]

        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if col == 2:
                cell.number_format = "0.00"

    widths = [30, 15, 40, 50, 50, 40]
    for col, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = width


def _write_listings_sheet(ws, ideas: List[Dict[str, Any]]):
    """Write listing templates to worksheet."""
    headers = [
        "Product Name", "Item Name", "Brand", "Product Type",
        "Bullet Points", "Description", "Material", "Colour", "Size",
        "SKU", "Price", "Condition", "Images"
    ]

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="BF8F00", end_color="BF8F00", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = thin_border

    for row_idx, idea in enumerate(ideas, 2):
        listing = idea.get("listing_template", {})
        identity = listing.get("product_identity", {})
        desc = listing.get("description", {})
        details = listing.get("product_details", {})
        offer = listing.get("offer", {})

        bullets = "\n".join([f"• {b}" for b in desc.get("bullet_points", []) if b])
        images = "\n".join(desc.get("images", [])[:3])

        data = [
            idea.get("name", ""),
            identity.get("item_name", ""),
            identity.get("brand_name", ""),
            identity.get("product_type", ""),
            bullets,
            desc.get("product_description", ""),
            details.get("material", ""),
            details.get("colour", ""),
            details.get("size", ""),
            offer.get("sku", ""),
            offer.get("your_price", ""),
            offer.get("item_condition", "New"),
            images,
        ]

        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row_idx, column=col, value=value)
            cell.border = thin_border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = [30, 30, 15, 20, 50, 50, 20, 20, 20, 15, 12, 12, 40]
    for col, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = width


def export_to_word(
    idea: Dict[str, Any],
    output_path: str,
    filename: Optional[str] = None,
) -> str:
    """Export a single product idea to Word document."""
    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    _add_title(doc, idea)

    _add_product_identity(doc, idea)

    _add_listing_content(doc, idea)

    _add_marketing_analysis(doc, idea)

    _add_metadata(doc, idea)

    if not filename:
        safe_name = "".join(c if c.isalnum() else "_" for c in idea.get("name", "product"))
        filename = f"{safe_name[:50]}.docx"

    filepath = os.path.join(output_path, filename)
    doc.save(filepath)
    return filepath


def _add_title(doc, idea: Dict[str, Any]):
    """Add document title."""
    title = doc.add_heading("Amazon Product Listing", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(idea.get("name", "Unknown Product"))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(47, 84, 150)

    priority = idea.get("priority", {})
    if priority:
        priority_para = doc.add_paragraph()
        priority_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tier = priority.get("tier", "unknown")
        rank = priority.get("rank", 0)
        run = priority_para.add_run(f"PRIORITY: {tier} (Rank #{rank})")
        run.font.size = Pt(12)
        run.bold = True
        if tier == "CRITICAL":
            run.font.color.rgb = RGBColor(192, 0, 0)
        elif tier == "HIGH":
            run.font.color.rgb = RGBColor(255, 165, 0)
        elif tier == "MEDIUM":
            run.font.color.rgb = RGBColor(0, 112, 192)

        action = priority.get("action", "")
        if action:
            action_para = doc.add_paragraph()
            action_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = action_para.add_run(f"Recommended Action: {action}")
            run.font.size = Pt(10)
            run.italic = True

    doc.add_paragraph()


def _add_product_identity(doc, idea: Dict[str, Any]):
    """Add product identity section."""
    doc.add_heading("Product Identity", level=1)

    listing = idea.get("listing_template", {})
    identity = listing.get("product_identity", {})
    table = doc.add_table(rows=6, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    rows_data = [
        ("ASIN", idea.get("asin", "N/A")),
        ("Category", idea.get("category", "N/A")),
        ("Brand", identity.get("brand_name", "(To be filled)")),
        ("Product Type", identity.get("product_type", "N/A")),
        ("Score", f"{idea.get('score', 0):.2f}"),
        ("Tier", idea.get("tier", "unknown")),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(10)

    doc.add_paragraph()


def _add_listing_content(doc, idea: Dict[str, Any]):
    """Add listing content section."""
    doc.add_heading("Listing Content", level=1)

    listing = idea.get("listing_template", {})
    desc = listing.get("description", {})

    doc.add_heading("Bullet Points", level=2)
    for bullet in desc.get("bullet_points", []):
        if bullet:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Product Description", level=2)
    description = desc.get("product_description", "")
    if description:
        doc.add_paragraph(description)
    else:
        doc.add_paragraph("(To be written - Focus on key features and benefits)").italic = True

    doc.add_heading("Images", level=2)
    images = desc.get("images", [])
    if images:
        for i, img_url in enumerate(images[:5], 1):
            doc.add_paragraph(f"Image {i}: {img_url}", style="List Number")
    else:
        doc.add_paragraph("(No images available)").italic = True

    doc.add_paragraph()


def _add_marketing_analysis(doc, idea: Dict[str, Any]):
    """Add marketing analysis section."""
    doc.add_heading("Marketing Analysis", level=1)

    marketing = idea.get("marketing", {})
    if not marketing:
        doc.add_paragraph("No marketing analysis available.")
        return

    summary = marketing.get("summary", "")
    if summary:
        p = doc.add_paragraph()
        run = p.add_run("Summary: ")
        run.bold = True
        p.add_run(summary)

    score = marketing.get("marketing_score", 0)
    p = doc.add_paragraph()
    run = p.add_run("Marketing Score: ")
    run.bold = True
    p.add_run(f"{score:.2f} / 1.00")

    doc.add_heading("Identified Problems", level=2)
    problems = marketing.get("problems", [])
    if problems:
        for problem in problems:
            severity = problem.get("severity", "").upper()
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{severity}] ")
            run.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0) if severity == "HIGH" else RGBColor(0, 112, 192)
            p.add_run(f"{problem.get('problem', '')}: ")
            p.add_run(problem.get("description", ""))
    else:
        doc.add_paragraph("No major problems identified.")

    doc.add_heading("Recommended Solutions", level=2)
    solutions = marketing.get("solutions", [])
    if solutions:
        for solution in solutions:
            doc.add_heading(f"{solution.get('solution', '')}", level=3)
            p = doc.add_paragraph()
            run = p.add_run("Priority: ")
            run.bold = True
            p.add_run(solution.get("priority", ""))

            p = doc.add_paragraph()
            run = p.add_run("Estimated Cost: ")
            run.bold = True
            p.add_run(solution.get("estimated_cost", "N/A"))

            p = doc.add_paragraph()
            run = p.add_run("Timeline: ")
            run.bold = True
            p.add_run(solution.get("timeline", "N/A"))

            doc.add_heading("Action Items:", level=3)
            for action in solution.get("actions", []):
                doc.add_paragraph(action, style="List Bullet")
    else:
        doc.add_paragraph("No specific solutions needed.")

    doc.add_heading("Recommended Marketing Strategies", level=2)
    strategies = marketing.get("recommended_strategies", [])
    if strategies:
        table = doc.add_table(rows=len(strategies) + 1, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        headers = ["Strategy", "Priority", "Cost", "Timeline"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            for paragraph in table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for i, strategy in enumerate(strategies, 1):
            table.rows[i].cells[0].text = strategy.get("name", "")
            table.rows[i].cells[1].text = strategy.get("priority", "")
            table.rows[i].cells[2].text = strategy.get("cost", "")
            table.rows[i].cells[3].text = strategy.get("time_to_results", "")

    doc.add_paragraph()


def _add_metadata(doc, idea: Dict[str, Any]):
    """Add metadata section."""
    doc.add_heading("Additional Information", level=1)

    listing = idea.get("listing_template", {})
    metadata = listing.get("_metadata", {})
    offer = listing.get("offer", {})
    details = listing.get("product_details", {})

    table = doc.add_table(rows=10, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    rows_data = [
        ("Price", f"£{offer.get('your_price', '0')}"),
        ("SKU", offer.get("sku", "N/A")),
        ("Condition", offer.get("item_condition", "New")),
        ("Material", details.get("material", "(To be filled)")),
        ("Colour", details.get("colour", "(To be filled)")),
        ("Size", details.get("size", "(To be filled)")),
        ("Estimated Margin", f"{metadata.get('estimated_margin_pct', 0):.0f}%"),
        ("Gated Category", "Yes - " + idea.get("gated_category", "") if idea.get("gated") else "No"),
        ("Source URL", idea.get("url", "N/A")),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]

    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True


def export_all_to_word(
    ideas: List[Dict[str, Any]],
    output_path: str,
) -> List[str]:
    """Export all product ideas to individual Word documents."""
    filepath = Path(output_path)
    filepath.mkdir(parents=True, exist_ok=True)

    exported = []
    for i, idea in enumerate(ideas, 1):
        safe_name = "".join(c if c.isalnum() else "_" for c in idea.get("name", "product"))
        filename = f"{i:03d}_{safe_name[:40]}.docx"
        path = export_to_word(idea, str(filepath), filename)
        exported.append(path)

    return exported
